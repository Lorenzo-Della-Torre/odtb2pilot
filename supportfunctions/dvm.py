"""
Module for creating DVMs from appropriately formatted yaml contained in the
module and function docstrings.
"""

import importlib.util
from pathlib import Path

import logging
import yaml
from html_to_docx import add_html
from markdown import markdown

from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

def get_reqdata(test_file_py):
    """get dict with all required reqdata attributes"""
    spec = importlib.util.spec_from_file_location("req_test", test_file_py)
    req_test = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(req_test)

    try:
        reqdata = yaml.safe_load(req_test.__doc__)
    except yaml.scanner.ScannerError:
        reqdata = {}

    reqdata.setdefault('reqprod', '')
    reqdata.setdefault('version', '')
    reqdata.setdefault('title', '')
    reqdata.setdefault('purpose', '')
    reqdata.setdefault('description', '')
    reqdata.setdefault('precondition', 'Default session')
    reqdata.setdefault('postcondition', 'Default session')
    reqdata.setdefault('details', '')

    dut_is_imported = False
    if 'Dut' in dir(req_test):
        dut_is_imported = True
        logging.debug("dut style testcase")

    return req_test, reqdata, dut_is_imported

#pylint: disable=too-many-statements
def create_dvm(test_file_py):
    """
    creates a docx document in the current work directory from test_file_py
    """
    req_test, reqdata, _ = get_reqdata(test_file_py)

    steps = []
    for step_name in dir(req_test):
        if step_name.startswith("step_"):
            step = getattr(req_test, step_name)
            doc = yaml.safe_load(step.__doc__)
            doc.setdefault('action', '')
            doc.setdefault('expected_result', '')
            doc.setdefault('result', '')
            doc.setdefault('comment', '')
            steps.append(doc)

    dvm = Document()

    section = dvm.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE

    # let's make the document A4 and not letter
    section = dvm.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    styles = dvm.styles
    heading = styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
    heading.font.name = 'Calibri Light'
    heading.font.size = Pt(28)

    body = styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = 'Calibri'
    body.font.size = Pt(11)
    body.font.bold = True

    normal_body = styles.add_style('Body Normal', WD_STYLE_TYPE.PARAGRAPH)
    normal_body.font.name = 'Calibri'
    normal_body.font.size = Pt(11)

    table_body = styles.add_style('Table Body', WD_STYLE_TYPE.TABLE)
    table_body.font.name = 'Calibri'
    table_body.font.size = Pt(11)

    border_table_body = styles['Table Grid']
    border_table_body.font.name = 'Calibri'
    border_table_body.font.size = Pt(11)

    dvm.styles.default(normal_body)

    dvm_heading = f'REQPROD {reqdata["reqprod"]} / MAIN ; {reqdata["version"]}'

    par = dvm.add_paragraph(dvm_heading)
    par.style = dvm.styles['Heading']

    par = dvm.add_paragraph('Tprocedure:')
    par.style = dvm.styles['Heading']

    records = (
        ('General:', dvm_heading),
        ('Title:', reqdata['title']),
        ('Purpose:', reqdata['purpose']),
    )

    table = dvm.add_table(rows=0, cols=2)
    table.style = table_body

    for title, content in records:
        row = table.add_row()
        par = row.cells[0].paragraphs[0]
        par.style = body
        par.text = title
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        par = row.cells[1].paragraphs[0]
        par.text = content

    row = table.add_row()
    par = row.cells[0].paragraphs[0]
    par.style = body
    par.text = 'Description:'
    par = row.cells[1].paragraphs[0]
    description = reqdata['description']
    for paragraph in description.split('\n'):
        html = markdown(paragraph)
        if html:
            par = add_html(par, html)

    table.columns[0].width = Mm(26)
    table.columns[1].width = Mm(218.4)

    table = dvm.add_table(rows=1, cols=2)
    table.style = table_body
    row = table.rows[0]
    par = row.cells[0].paragraphs[0]
    par.style = body
    par.text = "Precondition:"
    par = row.cells[1].paragraphs[0]
    par.text = reqdata["precondition"]
    table.columns[0].width = Mm(27)
    table.columns[1].width = Mm(216)

    dvm.add_page_break()

    par = dvm.add_paragraph("Test execution:")
    par.style = body

    table = dvm.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Step'
    hdr_cells[1].text = 'Action (client)'
    hdr_cells[2].text = 'Expected Result (ECU)'
    hdr_cells[3].text = 'Result'
    hdr_cells[4].text = 'Comment'

    for step, doc in enumerate(steps, 1):
        row = table.add_row()
        row.cells[0].text = str(step)
        row.cells[1].text = doc["action"]
        row.cells[2].text = doc["expected_result"]
        row.cells[3].text = doc["result"]
        row.cells[4].text = doc["comment"]

    table.columns[0].width = Mm(19)
    table.columns[1].width = Mm(71.3)
    table.columns[2].width = Mm(32.5)
    table.columns[3].width = Mm(14.8)
    table.columns[4].width = Mm(101.2)

    records = (
        ('Postcondition:', reqdata['postcondition']),
        ('Details:', reqdata['details']),
        ('Path:', '-'),
    )

    table = dvm.add_table(rows=0, cols=2)
    table.style = table_body

    for title, content in records:
        row = table.add_row()
        par = row.cells[0].paragraphs[0]
        par.style = body
        par.text = title
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].text = content

    table.columns[0].width = Mm(27)
    table.columns[1].width = Mm(216)

    docx_filename = Path(test_file_py).with_suffix('.docx').name
    docx_filename = docx_filename.replace("BSW_REQPROD_", "REQ_")
    print("Generated file: ", docx_filename)

    dvm.save(docx_filename)
